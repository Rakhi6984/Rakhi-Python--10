from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



# ============================================================
# PERSONAL INFORMATION - EDIT THESE DETAILS
# ============================================================

NAME = "Rakhi"
PHONE = "+91-9817865202"
EMAIL = "rakhi09jagadhri@gmail.com"
LOCATION = "Yamunanagar, Haryana, India"
LINKEDIN = "linkedin.com/in/Rakhi Chopra"
GITHUB = "github.com/rakhi09jagadhri"

CAREER_OBJECTIVE = (
    "Motivated BCA graduate seeking an entry-level Cloud/DevOps role "
    "where I can apply my knowledge of AWS Cloud, Python, Linux, and "
    "cloud automation while continuously developing my technical skills."
)


# ============================================================
# DOCUMENT SETUP
# ============================================================

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)


# Default font
styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10)

# Make font work properly in Word
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def add_bottom_border(paragraph):
    """
    Adds a horizontal line below a paragraph.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()

    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")

    pbdr.append(bottom)
    pPr.append(pbdr)


def add_heading(text):
    """
    Adds ATS-friendly section heading.
    """
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)

    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)

    add_bottom_border(paragraph)

    return paragraph


def add_bullet(text):
    """
    Adds a simple bullet point.
    """
    paragraph = doc.add_paragraph(style=None)

    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.12)
    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run("• " + text)
    run.font.name = "Arial"
    run.font.size = Pt(10)

    return paragraph


def add_normal(text):
    """
    Adds normal resume text.
    """
    paragraph = doc.add_paragraph()

    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)

    return paragraph


# ============================================================
# HEADER
# ============================================================

name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_para.paragraph_format.space_after = Pt(2)

name_run = name_para.add_run(NAME)
name_run.bold = True
name_run.font.name = "Arial"
name_run.font.size = Pt(20)


contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_para.paragraph_format.space_after = Pt(6)

contact_run = contact_para.add_run(
    f"{PHONE} | {EMAIL} | {LOCATION}\n"
    f"{LINKEDIN} | {GITHUB}"
)

contact_run.font.name = "Arial"
contact_run.font.size = Pt(9)


# ============================================================
# CAREER OBJECTIVE
# ============================================================

add_heading("Career Objective")

add_normal(CAREER_OBJECTIVE)


# ============================================================
# EDUCATION
# ============================================================

add_heading("Education")

education = doc.add_paragraph()
education.paragraph_format.space_after = Pt(2)

run = education.add_run("Bachelor of Computer Applications (BCA)")
run.bold = True
run.font.size = Pt(10)

education.add_run(
    "\nS.D. Institute of Management and Technology, Jagadhri"
    "\nKurukshetra University"
    "\nExpected Graduation: 2026"
)


# ============================================================
# TECHNICAL SKILLS
# ============================================================

add_heading("Technical Skills")

skills = [
    ("Cloud Platform", "Amazon Web Services (AWS)"),
    ("AWS Services", "EC2, S3, IAM, VPC, CloudWatch, Lambda"),
    ("Programming", "Python"),
    ("Cloud Automation", "AWS CLI, Boto3"),
    ("Operating Systems", "Linux, Windows"),
    ("Version Control", "Git, GitHub"),
    ("Containers", "Docker"),
    ("Web Server", "Apache HTTP Server"),
    ("Database", "MySQL, SQL"),
    ("Tools", "VS Code, WSL, AWS Management Console"),
]


for category, technologies in skills:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)

    r1 = paragraph.add_run(category + ": ")
    r1.bold = True
    r1.font.size = Pt(10)

    r2 = paragraph.add_run(technologies)
    r2.font.size = Pt(10)


# ============================================================
# PROJECTS
# ============================================================

add_heading("Projects")


# Project 1
project = doc.add_paragraph()
project.paragraph_format.space_after = Pt(2)

r = project.add_run("EC2 and S3 Cloud Automation using Python")
r.bold = True
r.font.size = Pt(10)

add_bullet(
    "Developed Python scripts using Boto3 to automate AWS EC2 and Amazon S3 operations."
)

add_bullet(
    "Automated cloud resource tasks such as EC2 instance management and S3 bucket operations."
)

add_bullet(
    "Used AWS IAM credentials and region configuration to securely interact with AWS services."
)

add_bullet(
    "Designed the project to demonstrate practical cloud automation using Python and AWS."
)


# Project 2
project = doc.add_paragraph()
project.paragraph_format.space_after = Pt(2)

r = project.add_run("Dockerized Web Application")
r.bold = True
r.font.size = Pt(10)

add_bullet(
    "Created a Docker-based web application environment using a Dockerfile."
)

add_bullet(
    "Built and managed Docker images and containers using Docker CLI commands."
)

add_bullet(
    "Configured application port mapping and practiced container lifecycle management."
)

add_bullet(
    "Deployed and tested the application in a Linux-based environment."
)


# Project 3
project = doc.add_paragraph()
project.paragraph_format.space_after = Pt(2)

r = project.add_run("AWS EC2 Web Server Deployment")
r.bold = True
r.font.size = Pt(10)

add_bullet(
    "Launched and configured an AWS EC2 instance for hosting a web server."
)

add_bullet(
    "Configured Apache HTTP Server and deployed an HTML web page."
)

add_bullet(
    "Worked with Linux commands, SSH access, security groups, and server configuration."
)


# ============================================================
# INTERNSHIP / TRAINING
# ============================================================

add_heading("Training / Internship")

training = doc.add_paragraph()
training.paragraph_format.space_after = Pt(2)

r = training.add_run("DevOps and Cloud Computing Training")
r.bold = True
r.font.size = Pt(10)

add_bullet(
    "Gained practical exposure to AWS Cloud, Linux, Python, Docker, Git, and cloud automation."
)

add_bullet(
    "Worked on AWS EC2 and S3 automation using Python and Boto3."
)

add_bullet(
    "Practiced deploying applications and managing cloud infrastructure."
)

add_bullet(
    "Worked with AWS services including EC2, S3, IAM, VPC, DynamoDB, and CloudWatch."
)


# ============================================================
# CERTIFICATIONS
# ============================================================

add_heading("Certifications & Training")

add_bullet("Python with AWS Cloud – Project-Based Training")
add_bullet("AWS Cloud Fundamentals and Boto3 Automation Training")
add_bullet("DevOps and Cloud Computing Internship/Training")


# ============================================================
# KEY STRENGTHS
# ============================================================

add_heading("Key Strengths")

add_bullet("Strong interest in Cloud Computing and DevOps")
add_bullet("Problem-solving and troubleshooting")
add_bullet("Quick learner with hands-on technical approach")
add_bullet("Good understanding of Linux and cloud environments")
add_bullet("Ability to learn and work with new technologies")


# ============================================================
# DECLARATION
# ============================================================

add_heading("Declaration")

add_normal(
    "I hereby declare that the information provided above is true and correct "
    "to the best of my knowledge and belief."
)


# ============================================================
# SAVE DOCUMENT
# ============================================================

output_file = "Cloud_AWS_Fresher_Resume.docx"

doc.save(output_file)

print("=" * 55)
print("RESUME CREATED SUCCESSFULLY")
print("=" * 55)
print(f"File: {output_file}")
print("Open the generated DOCX file in Microsoft Word.")
print("=" * 55)